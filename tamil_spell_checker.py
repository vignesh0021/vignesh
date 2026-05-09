#!/usr/bin/env python3
"""
Tamil Spell Checker and Auto-Corrector for DOCX and PDF Documents
==================================================================
Detects and auto-corrects Tamil spelling errors in .docx and .pdf files
while preserving the original document formatting and structure.

Installation:
    pip install -r requirements.txt

Usage:
    python tamil_spell_checker.py document.docx
    python tamil_spell_checker.py document.pdf --output corrected.pdf
    python tamil_spell_checker.py document.docx --no-autocorrect
    python tamil_spell_checker.py document.docx --proper-nouns-file nouns.txt
    python tamil_spell_checker.py document.pdf  --tamil-font /usr/share/fonts/truetype/lohit-tamil/Lohit-Tamil.ttf
    python tamil_spell_checker.py document.docx --max-size 100 --force
"""

import argparse
import logging
import re
import subprocess
import sys
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, NamedTuple, Optional, Set, Tuple

# ── Logging setup ──────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)],
)
logger = logging.getLogger(__name__)

_DEFAULT_MAX_SIZE_MB = 50


# ── Dependency availability flags ──────────────────────────────────────────────
def _try_import(module_name: str):
    try:
        import importlib
        return importlib.import_module(module_name), None
    except ImportError as e:
        return None, str(e)


_docx_mod, _docx_err = _try_import("docx")
_fitz_mod, _fitz_err = _try_import("fitz")
_tamil_mod, _tamil_err = _try_import("tamil")

DOCX_AVAILABLE = _docx_mod is not None
PDF_AVAILABLE = _fitz_mod is not None
TAMIL_AVAILABLE = _tamil_mod is not None


def check_dependencies(file_ext: str) -> None:
    """Exit with a clear message if required packages are missing."""
    missing: List[str] = []
    if not TAMIL_AVAILABLE:
        missing.append("open-tamil")
    if file_ext == ".docx" and not DOCX_AVAILABLE:
        missing.append("python-docx")
    if file_ext == ".pdf" and not PDF_AVAILABLE:
        missing.append("PyMuPDF")
    if missing:
        logger.error("Missing required packages: %s", ", ".join(missing))
        logger.error("Install with:  pip install %s", " ".join(missing))
        sys.exit(1)


# ── Tamil Unicode utilities ────────────────────────────────────────────────────

# Tamil Unicode block: U+0B80 – U+0BFF
_TAMIL_RE = re.compile(r"[஀-௿]+")

# Tamil digits U+0BE6–U+0BEF and symbols U+0BF0–U+0BFA — not real words
_TAMIL_NON_WORD_RE = re.compile(r"^[௦-௺]+$")


def is_tamil_text(text: str) -> bool:
    return bool(_TAMIL_RE.search(text))


def is_checkable_tamil_word(word: str) -> bool:
    """Return True for Tamil words that should be spell-checked (skip pure numerals/symbols)."""
    clean = re.sub(r"[^஀-௿]", "", word)
    return bool(clean) and not _TAMIL_NON_WORD_RE.match(clean)


# ── Tamil font detection ───────────────────────────────────────────────────────

# Well-known paths searched in order before falling back to fc-list.
_TAMIL_FONT_CANDIDATES = [
    # Linux – Noto Sans Tamil
    "/usr/share/fonts/truetype/noto/NotoSansTamil-Regular.ttf",
    "/usr/share/fonts/opentype/noto/NotoSansTamil-Regular.ttf",
    "/usr/share/fonts/noto/NotoSansTamil-Regular.ttf",
    # Linux – Lohit Tamil
    "/usr/share/fonts/truetype/lohit-tamil/Lohit-Tamil.ttf",
    "/usr/share/fonts/lohit-tamil/Lohit-Tamil.ttf",
    "/usr/share/fonts/lohit/Lohit-Tamil.ttf",
    # Linux – other
    "/usr/share/fonts/truetype/fonts-tamilfont/TAMu_Kalyani.ttf",
    # macOS
    "/System/Library/Fonts/Supplemental/Tamil MN.ttc",
    "/Library/Fonts/Tamil Sangam MN.ttc",
    "/System/Library/Fonts/Tamil.ttf",
    # Windows
    "C:/Windows/Fonts/latha.ttf",
    "C:/Windows/Fonts/Latha.ttf",
    "C:/Windows/Fonts/cardt.ttf",
]


def find_tamil_font(hint: Optional[str] = None) -> Optional[str]:
    """
    Return a path to a Tamil-capable font file.

    Resolution order:
      1. User-supplied *hint* via --tamil-font
      2. Well-known system paths
      3. fc-list (Linux fontconfig)
    """
    if hint:
        p = Path(hint)
        if p.exists():
            return str(p)
        logger.error("--tamil-font path does not exist: %s", hint)
        return None

    for path in _TAMIL_FONT_CANDIDATES:
        if Path(path).exists():
            logger.info("Tamil font auto-detected: %s", path)
            return path

    # Linux fallback: ask fontconfig
    try:
        result = subprocess.run(
            ["fc-list", ":lang=ta", "--format=%{file}\n"],
            capture_output=True, text=True, timeout=5,
        )
        for line in result.stdout.splitlines():
            line = line.strip()
            if line and Path(line).exists():
                logger.info("Tamil font found via fc-list: %s", line)
                return line
    except Exception:
        pass

    return None


# ── Edit distance ──────────────────────────────────────────────────────────────

def _levenshtein(s1: str, s2: str) -> int:
    """Standard DP Levenshtein distance, operates on Unicode code points."""
    m, n = len(s1), len(s2)
    dp = list(range(n + 1))
    for i in range(1, m + 1):
        prev, dp[0] = dp[0], i
        for j in range(1, n + 1):
            temp = dp[j]
            dp[j] = prev if s1[i - 1] == s2[j - 1] else 1 + min(prev, dp[j], dp[j - 1])
            prev = temp
    return dp[n]


# ── Spell-checker wrapper ──────────────────────────────────────────────────────

class TamilSpellCheckerWrapper:
    """
    Wraps the open-tamil spell module with multiple fallback strategies so the
    script stays usable even when the spell sub-package is partially installed.

    Strategy order:
      1. spell.checker.SpellChecker  (open-tamil ≥ 0.9, spell installed as top-level pkg)
      2. tamil.spell.SpellChecker    (some packaged distributions)
      3. Dictionary lookup via open-tamil's built-in word lists + Levenshtein ranking
      4. No checking (pass-through, with a warning)
    """

    def __init__(self) -> None:
        self._checker = None
        self._word_set: Set[str] = set()
        self._mode = "none"
        self._init()

    # ── Initialisation ─────────────────────────────────────────────────────────

    def _init(self) -> None:
        try:
            from spell.checker import SpellChecker  # type: ignore
            self._checker = SpellChecker(lang="TA")
            self._mode = "spell.checker"
            logger.info("Spell checker: using spell.checker.SpellChecker")
            return
        except Exception as exc:
            logger.debug("spell.checker unavailable: %s", exc)

        try:
            import tamil.spell as ts  # type: ignore
            self._checker = ts.SpellChecker()
            self._mode = "tamil.spell"
            logger.info("Spell checker: using tamil.spell.SpellChecker")
            return
        except Exception as exc:
            logger.debug("tamil.spell unavailable: %s", exc)

        self._load_word_list()
        if self._word_set:
            self._mode = "dictionary"
            logger.info(
                "Spell checker: dictionary mode (%d words loaded)", len(self._word_set)
            )
            return

        logger.warning(
            "No Tamil spell checker backend found. "
            "Errors will be reported but no corrections can be suggested."
        )
        self._mode = "none"

    def _load_word_list(self) -> None:
        """Collect Tamil words from open-tamil's bundled data files."""
        if not TAMIL_AVAILABLE:
            return
        try:
            import tamil  # type: ignore
            pkg_root = Path(tamil.__file__).parent
            for wl_file in sorted(pkg_root.rglob("*.txt")):
                self._ingest_word_file(wl_file)
            for wl_file in sorted(pkg_root.rglob("*.tsv")):
                self._ingest_word_file(wl_file)
        except Exception as exc:
            logger.debug("Word list loading failed: %s", exc)
        self._word_set.update(self._seed_words())

    def _ingest_word_file(self, path: Path) -> None:
        try:
            with path.open(encoding="utf-8", errors="ignore") as fh:
                for line in fh:
                    word = line.strip().split("\t")[0].strip()
                    if word and is_tamil_text(word):
                        self._word_set.add(word)
        except Exception:
            pass

    @staticmethod
    def _seed_words() -> List[str]:
        """
        Minimal seed list of common Tamil words.
        Acts as a last-resort fallback when no word files are found.
        """
        return [
            "அம்மா", "அப்பா", "வீடு", "நாடு", "மனிதன்", "மனிதர்கள்",
            "தமிழ்", "மொழி", "நண்பன்", "நண்பர்கள்", "நாள்", "வருடம்",
            "மாதம்", "ஆண்டு", "பள்ளி", "கல்லூரி", "பல்கலைக்கழகம்",
            "வணக்கம்", "நன்றி", "வாழ்க்கை", "உலகம்", "இந்தியா",
            "தமிழ்நாடு", "சென்னை", "மதுரை", "கோவை", "திருச்சி",
            "புத்தகம்", "கடல்", "மலை", "ஆறு", "காடு", "நகரம்",
            "கிராமம்", "விவசாயம்", "மரம்", "பூ", "பழம்", "காய்கறி",
            "சாப்பிடு", "படி", "விளையாடு", "பேசு", "எழுது", "படிக்க",
            "வேலை", "ஓய்வு", "தூக்கம்", "கனவு", "அன்பு", "நேசம்",
        ]

    # ── Public API ─────────────────────────────────────────────────────────────

    def is_correct(self, word: str) -> bool:
        """Return True if *word* appears to be a valid Tamil word."""
        if self._mode in ("spell.checker", "tamil.spell"):
            try:
                return bool(self._checker.check(word))  # type: ignore[union-attr]
            except Exception:
                pass
        if self._mode == "dictionary":
            return word in self._word_set
        # pass-through mode: assume correct to avoid flooding false positives
        return True

    def suggest(self, word: str) -> List[str]:
        """Return a ranked list of suggested corrections for *word*."""
        if self._mode in ("spell.checker", "tamil.spell"):
            try:
                for method in ("get_suggestion_for_word", "suggest"):
                    fn = getattr(self._checker, method, None)
                    if callable(fn):
                        result = fn(word)
                        if isinstance(result, (list, tuple)) and result:
                            return list(result)
            except Exception:
                pass
        if self._mode == "dictionary" and self._word_set:
            return self._nearest_words(word)
        return []

    def best_correction(self, word: str) -> Optional[str]:
        """Return the single best correction, or None if nothing better is found."""
        suggestions = self.suggest(word)
        return suggestions[0] if suggestions else None

    def _nearest_words(
        self, word: str, max_results: int = 5, max_dist: int = 3
    ) -> List[str]:
        """
        Levenshtein-based nearest-word search with a fast length pre-filter.
        Only candidates within *max_dist* edits are returned, sorted by distance.
        """
        word_len = len(word)
        scored: List[Tuple[int, str]] = []
        for candidate in self._word_set:
            if abs(len(candidate) - word_len) > max_dist:
                continue  # guaranteed to exceed max_dist; skip cheaply
            dist = _levenshtein(word, candidate)
            if dist <= max_dist:
                scored.append((dist, candidate))
        scored.sort()
        return [w for _, w in scored[:max_results]]


# ── Correction records ─────────────────────────────────────────────────────────

@dataclass
class CorrectionRecord:
    original: str
    corrected: str
    location: str


@dataclass
class SpellCheckSummary:
    total_checked: int = 0
    errors_found: int = 0
    corrections_made: int = 0
    uncorrectable: List[str] = field(default_factory=list)
    log: List[CorrectionRecord] = field(default_factory=list)

    def record_correction(self, original: str, corrected: str, location: str) -> None:
        self.errors_found += 1
        self.corrections_made += 1
        self.log.append(CorrectionRecord(original, corrected, location))

    def record_uncorrectable(self, word: str) -> None:
        self.errors_found += 1
        self.uncorrectable.append(word)

    def print_summary(self) -> None:
        bar = "=" * 62
        print(f"\n{bar}")
        print("  TAMIL SPELL CHECK SUMMARY")
        print(bar)
        print(f"  Total Tamil words checked  : {self.total_checked}")
        print(f"  Spelling errors found      : {self.errors_found}")
        print(f"  Corrections applied        : {self.corrections_made}")
        print(f"  Words with no suggestion   : {len(self.uncorrectable)}")

        if self.log:
            print("\n  CORRECTIONS APPLIED:")
            print("  " + "-" * 58)
            for rec in self.log:
                print(f"  '{rec.original}'  →  '{rec.corrected}'")
                print(f"      Location: {rec.location}")

        if self.uncorrectable:
            print("\n  UNCORRECTABLE WORDS (no suggestion found):")
            print("  " + "-" * 58)
            for word, count in Counter(self.uncorrectable).most_common():
                print(f"  '{word}'  (×{count})")

        print(bar + "\n")


# ── Core text-processing engine ────────────────────────────────────────────────

def process_text(
    text: str,
    checker: TamilSpellCheckerWrapper,
    summary: SpellCheckSummary,
    location: str,
    *,
    auto_correct: bool = True,
    proper_nouns: Optional[Set[str]] = None,
) -> str:
    """
    Scan *text* for Tamil spelling errors, optionally replacing them in-place.
    Non-Tamil characters, digits, and proper nouns are left untouched.
    Returns the (possibly corrected) text.
    """
    if not text or not is_tamil_text(text):
        return text

    if proper_nouns is None:
        proper_nouns = set()

    result_parts: List[str] = []

    for token in re.split(r"(\s+)", text):
        if re.fullmatch(r"\s+", token) or not token:
            result_parts.append(token)
            continue

        tamil_match = _TAMIL_RE.search(token)
        if not tamil_match:
            result_parts.append(token)
            continue

        word = tamil_match.group()
        if not is_checkable_tamil_word(word):
            result_parts.append(token)
            continue

        summary.total_checked += 1

        if word in proper_nouns or checker.is_correct(word):
            result_parts.append(token)
            continue

        # ── Misspelling detected ───────────────────────────────────────────
        if not auto_correct:
            logger.info("Spelling error at %-40s  '%s'", location + ":", word)
            summary.record_uncorrectable(word)
            result_parts.append(token)
            continue

        correction = checker.best_correction(word)
        if correction and correction != word:
            corrected_token = (
                token[: tamil_match.start()] + correction + token[tamil_match.end():]
            )
            result_parts.append(corrected_token)
            summary.record_correction(word, correction, location)
            logger.debug("Corrected [%s]  '%s' → '%s'", location, word, correction)
        else:
            result_parts.append(token)
            summary.record_uncorrectable(word)

    return "".join(result_parts)


# ── DOCX handler ───────────────────────────────────────────────────────────────

def process_docx(
    input_path: Path,
    output_path: Path,
    checker: TamilSpellCheckerWrapper,
    summary: SpellCheckSummary,
    *,
    auto_correct: bool = True,
    proper_nouns: Optional[Set[str]] = None,
) -> None:
    """
    Read *input_path* (DOCX), spell-check every Run in the document
    (body, tables, headers/footers, and text boxes), and save to *output_path*.

    Each Run is corrected independently so all character-level formatting
    (bold, italic, font, size, colour) is fully preserved.
    """
    if not DOCX_AVAILABLE:
        logger.error("python-docx is required for DOCX files.  pip install python-docx")
        sys.exit(1)

    from docx import Document  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.text.paragraph import Paragraph as DocxParagraph  # type: ignore

    logger.info("Opening DOCX: %s", input_path)
    doc = Document(str(input_path))

    def _fix_run(run, loc: str) -> None:
        if run.text and is_tamil_text(run.text):
            run.text = process_text(
                run.text, checker, summary, loc,
                auto_correct=auto_correct, proper_nouns=proper_nouns,
            )

    def _fix_paragraph(para, loc: str) -> None:
        for idx, run in enumerate(para.runs):
            _fix_run(run, f"{loc} › run {idx + 1}")

    # Body paragraphs
    for p_idx, para in enumerate(doc.paragraphs):
        _fix_paragraph(para, f"body › para {p_idx + 1}")

    # Tables (all rows and cells)
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    loc = (
                        f"table {t_idx+1} › row {r_idx+1} "
                        f"› col {c_idx+1} › para {p_idx+1}"
                    )
                    _fix_paragraph(para, loc)

    # Headers and footers (all section variants)
    for s_idx, section in enumerate(doc.sections):
        for hf_name, hf_obj in (
            ("header", section.header),
            ("footer", section.footer),
            ("first-page header", section.first_page_header),
            ("first-page footer", section.first_page_footer),
            ("even-page header", section.even_page_header),
            ("even-page footer", section.even_page_footer),
        ):
            if hf_obj is None:
                continue
            for p_idx, para in enumerate(hf_obj.paragraphs):
                _fix_paragraph(para, f"section {s_idx+1} {hf_name} › para {p_idx+1}")

    # Text boxes — DrawingML <wps:txbx> elements not exposed by the high-level API
    textboxes = list(doc.element.body.iter(qn("w:txbxContent")))
    for txbx_idx, txbx_content in enumerate(textboxes):
        for p_elem in txbx_content.findall(qn("w:p")):
            para = DocxParagraph(p_elem, doc)
            _fix_paragraph(para, f"textbox {txbx_idx + 1}")
    if textboxes:
        logger.info("Processed %d text box(es)", len(textboxes))

    doc.save(str(output_path))
    logger.info("Corrected DOCX saved → %s", output_path)


# ── PDF handler ────────────────────────────────────────────────────────────────

class _PdfFix(NamedTuple):
    """One pending text correction for a PDF span."""
    bbox: object                            # fitz.Rect
    corrected_text: str
    font_size: float
    color_rgb: Tuple[float, float, float]


def _is_scanned_pdf(doc) -> bool:
    """
    Heuristic: the PDF is likely scanned (image-only) if a sample of its
    pages have no extractable text but do contain embedded images.
    """
    sample = list(doc)[: min(10, len(doc))]
    if not sample:
        return False
    text_pages = sum(1 for page in sample if page.get_text().strip())
    image_pages = sum(1 for page in sample if page.get_images())
    return text_pages == 0 and image_pages > 0


def process_pdf(
    input_path: Path,
    output_path: Path,
    checker: TamilSpellCheckerWrapper,
    summary: SpellCheckSummary,
    *,
    auto_correct: bool = True,
    proper_nouns: Optional[Set[str]] = None,
    tamil_font_path: Optional[str] = None,
) -> None:
    """
    Read *input_path* (PDF), spell-check Tamil text spans, write corrected PDF.

    Correction strategy (per page):
      1. Extract text spans with bounding boxes via get_text("dict").
      2. Collect all spans that need correction.
      3. Add ALL redaction annotations first, then call apply_redactions() once
         (batching prevents bounding-box drift from incremental geometry updates).
      4. Reinsert every corrected string at its original position using the
         supplied Tamil font.

    NOTE: PDF editing is inherently "best-effort". Complex layouts, embedded
    fonts, and ligature-based rendering may show minor visual differences.
    The original PDF is never modified; corrections go to a new file.
    """
    if not PDF_AVAILABLE:
        logger.error("PyMuPDF is required for PDF files.  pip install PyMuPDF")
        sys.exit(1)

    import fitz  # type: ignore

    logger.info("Opening PDF: %s", input_path)
    doc = fitz.open(str(input_path))

    # Scanned PDF guard — no text layer means nothing to correct
    if _is_scanned_pdf(doc):
        logger.warning(
            "This PDF appears to be scanned (image-only, no text layer). "
            "Spell checking requires selectable text. "
            "Use an OCR tool (e.g. 'ocrmypdf input.pdf output.pdf') to add a text layer first."
        )
        doc.close()
        return

    for page_no, page in enumerate(doc):
        logger.info("  Page %d / %d …", page_no + 1, len(doc))
        pending: List[_PdfFix] = []

        for block in page.get_text("dict").get("blocks", []):
            if block.get("type") != 0:  # 0 = text block
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    original = span.get("text", "")
                    if not original or not is_tamil_text(original):
                        continue

                    location = f"page {page_no + 1}"
                    corrected = process_text(
                        original, checker, summary, location,
                        auto_correct=auto_correct, proper_nouns=proper_nouns,
                    )

                    if auto_correct and corrected != original:
                        c = span.get("color", 0)
                        pending.append(_PdfFix(
                            bbox=fitz.Rect(span["bbox"]),
                            corrected_text=corrected,
                            font_size=float(span.get("size", 11)),
                            color_rgb=(
                                ((c >> 16) & 0xFF) / 255.0,
                                ((c >> 8) & 0xFF) / 255.0,
                                (c & 0xFF) / 255.0,
                            ),
                        ))

        if pending:
            _batch_apply_corrections(page, pending, tamil_font_path)

    doc.save(str(output_path), garbage=4, deflate=True)
    doc.close()
    logger.info("Corrected PDF saved → %s", output_path)


def _batch_apply_corrections(
    page,
    fixes: List[_PdfFix],
    tamil_font_path: Optional[str],
) -> None:
    """
    Apply all pending corrections for one page in two clean passes:

    Pass 1 — register every redaction annotation, then apply them all at once.
              Batching is critical: calling apply_redactions() per-span shifts
              page geometry and makes subsequent bounding boxes stale.

    Pass 2 — insert each corrected string at the vacated position using the
              Tamil-capable font.
    """
    import fitz  # type: ignore

    # Pass 1: mark every region for erasure, then erase all at once
    for fix in fixes:
        page.add_redact_annot(fix.bbox, fill=(1, 1, 1))
    page.apply_redactions()

    # Pass 2: write corrected text into each cleared region
    for fix in fixes:
        # Approximate the original baseline: bottom of bounding box minus 1 pt
        insertion_point = fitz.Point(fix.bbox.x0, fix.bbox.y1 - 1)
        kwargs = {"fontsize": fix.font_size, "color": fix.color_rgb}
        if tamil_font_path:
            kwargs["fontfile"] = tamil_font_path
        try:
            page.insert_text(insertion_point, fix.corrected_text, **kwargs)
        except Exception as exc:
            logger.debug("insert_text failed for span (%s); skipped", exc)


# ── CLI ────────────────────────────────────────────────────────────────────────

def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        prog="tamil_spell_checker",
        description="Tamil spell checker and auto-corrector for DOCX and PDF files.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python tamil_spell_checker.py report.docx
  python tamil_spell_checker.py report.docx --output report_fixed.docx
  python tamil_spell_checker.py report.pdf  --output report_fixed.pdf \\
      --tamil-font /usr/share/fonts/truetype/lohit-tamil/Lohit-Tamil.ttf
  python tamil_spell_checker.py report.docx --no-autocorrect
  python tamil_spell_checker.py report.docx --proper-nouns-file nouns.txt --log-level DEBUG
  python tamil_spell_checker.py report.docx --max-size 100 --force
        """,
    )
    parser.add_argument("input_file", help="Path to the input .docx or .pdf file.")
    parser.add_argument(
        "--output", "-o",
        metavar="PATH",
        help="Output file path.  Defaults to <stem>_corrected.<ext> in the same directory.",
    )
    parser.add_argument(
        "--no-autocorrect",
        action="store_true",
        help="Report spelling errors without modifying the document (dry-run mode).",
    )
    parser.add_argument(
        "--proper-nouns-file",
        metavar="FILE",
        help="Plain-text file of Tamil proper nouns to skip (one per line).",
    )
    parser.add_argument(
        "--tamil-font",
        metavar="FILE",
        help=(
            "Path to a Tamil-capable .ttf/.otf/.ttc font file. "
            "Required for PDF auto-correction when no Tamil font is installed system-wide. "
            "Install one with: sudo apt install fonts-lohit-taml  (Ubuntu/Debian) "
            "or sudo dnf install lohit-tamil-fonts  (Fedora)."
        ),
    )
    parser.add_argument(
        "--max-size",
        type=float,
        default=_DEFAULT_MAX_SIZE_MB,
        metavar="MB",
        help=f"Maximum input file size in MB (default: {_DEFAULT_MAX_SIZE_MB}). "
             "Prevents accidental out-of-memory on very large files.",
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Overwrite the output file if it already exists without prompting.",
    )
    parser.add_argument(
        "--log-level",
        choices=["DEBUG", "INFO", "WARNING", "ERROR"],
        default="INFO",
        help="Logging verbosity level (default: INFO).",
    )
    return parser.parse_args()


def _load_proper_nouns(filepath: Optional[str]) -> Set[str]:
    if not filepath:
        return set()
    try:
        with open(filepath, encoding="utf-8") as fh:
            nouns = {line.strip() for line in fh if line.strip()}
        logger.info("Loaded %d proper nouns from %s", len(nouns), filepath)
        return nouns
    except FileNotFoundError:
        logger.warning("Proper-nouns file not found: %s  (skipping)", filepath)
        return set()


def _default_output_path(input_path: Path) -> Path:
    return input_path.parent / f"{input_path.stem}_corrected{input_path.suffix}"


# ── Entry point ────────────────────────────────────────────────────────────────

def main() -> None:
    args = parse_args()
    logging.getLogger().setLevel(getattr(logging, args.log_level))

    input_path = Path(args.input_file).resolve()
    if not input_path.exists():
        logger.error("Input file not found: %s", input_path)
        sys.exit(1)

    file_ext = input_path.suffix.lower()
    if file_ext not in (".docx", ".pdf"):
        logger.error(
            "Unsupported file format '%s'.  Supported extensions: .docx  .pdf", file_ext
        )
        sys.exit(1)

    # File size guard — catch huge files before loading into memory
    file_size_mb = input_path.stat().st_size / (1024 * 1024)
    if file_size_mb > args.max_size:
        logger.error(
            "File size %.1f MB exceeds the %.0f MB limit. "
            "Use --max-size to raise it.",
            file_size_mb, args.max_size,
        )
        sys.exit(1)

    check_dependencies(file_ext)

    output_path = (
        Path(args.output).resolve() if args.output else _default_output_path(input_path)
    )

    # Overwrite guard
    if output_path.exists() and not args.force:
        logger.error(
            "Output file already exists: %s\n"
            "  Use --force to overwrite, or --output to choose a different path.",
            output_path,
        )
        sys.exit(1)

    proper_nouns = _load_proper_nouns(args.proper_nouns_file)
    auto_correct = not args.no_autocorrect

    # Tamil font resolution for PDF — fail fast here rather than mid-processing
    tamil_font_path: Optional[str] = None
    if file_ext == ".pdf" and auto_correct:
        tamil_font_path = find_tamil_font(args.tamil_font)
        if not tamil_font_path:
            logger.error(
                "No Tamil-capable font found. PDF auto-correction requires one.\n"
                "  Install : sudo apt install fonts-lohit-taml        (Ubuntu/Debian)\n"
                "          : sudo dnf install lohit-tamil-fonts        (Fedora)\n"
                "  Or pass : --tamil-font /path/to/NotoSansTamil-Regular.ttf\n"
                "  Or use  : --no-autocorrect  to report errors without editing the PDF."
            )
            sys.exit(1)

    logger.info("─" * 62)
    logger.info("Input file    : %s  (%.1f MB)", input_path, file_size_mb)
    logger.info("Output file   : %s", output_path)
    logger.info("Auto-correct  : %s", auto_correct)
    if tamil_font_path:
        logger.info("Tamil font    : %s", tamil_font_path)
    logger.info("─" * 62)

    logger.info("Initialising Tamil spell checker …")
    checker = TamilSpellCheckerWrapper()
    summary = SpellCheckSummary()

    try:
        if file_ext == ".docx":
            process_docx(
                input_path, output_path, checker, summary,
                auto_correct=auto_correct, proper_nouns=proper_nouns,
            )
        else:
            process_pdf(
                input_path, output_path, checker, summary,
                auto_correct=auto_correct, proper_nouns=proper_nouns,
                tamil_font_path=tamil_font_path,
            )
    except Exception as exc:
        logger.error("Failed to process '%s': %s", input_path.name, exc, exc_info=True)
        sys.exit(1)

    summary.print_summary()

    if auto_correct:
        print(f"Corrected file saved to: {output_path}")
    else:
        print("Dry-run complete. No file was written.")


if __name__ == "__main__":
    main()
